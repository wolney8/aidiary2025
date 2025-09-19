# daily_file_ingestion.py

import glob
import sqlite3
import os
import re
from nltk.corpus import stopwords # type: ignore
from nltk.tokenize import WordPunctTokenizer # type: ignore
from collections import Counter
from docx import Document # type: ignore
import json
import nltk # type: ignore

# MODELS Import
from app.models.models import User, DailyDiaryEntry, db

# UTILS Import
from utils import get_continuous_chunks, process_text_for_nltk



# Get the path to the app's root directory
app_root_path = os.path.dirname(os.path.abspath(__file__))

# Set the NLTK data path to the nltk_data folder in the app's root directory
nltk.data.path.append(os.path.join(app_root_path, 'nltk_data'))


# Read stop words from JSON file
with open('stop_words.json', 'r') as file:
    custom_stop_words = json.load(file)["stop_words"]


# # # # # # 
# Database connection logic
# # # # # # 

# Inserts an entry into the daily diary database.
# New function using SQLAlchemy
def insert_into_db_ingestion(user_id, entry_date, entry_number, user_message, ai_response, daily_people_names, tags_str):
    new_entry = DailyDiaryEntry(
        user_id=user_id,
        entry_date=entry_date,
        entry_number=entry_number,  # Include the entry_number parameter
        user_message=user_message,
        ai_response=ai_response,
        daily_people_names=daily_people_names,
        tags=tags_str
    )
    db.session.add(new_entry)
    db.session.commit()



# # # # # # 
# FILE UPLOAD Routes and Endpoints
# DAILY DIARY
# # # # # # 
# Reads a Word document containing a user's message and ChatGPT's response. It combines the text and processes it using similar techniques as in the Dream Diary (tokenization, stopword removal, part-of-speech tagging, frequency analysis) to generate tags. The extracted and processed information, including the tags, is then inserted into the daily diary database.

def ingest_word_doc(user_id, file_path, entry_date):
    doc = Document(file_path)
    entries = []
    current_entry = {"user_message": "", "ai_response": ""}

    # Join all paragraphs together to handle line breaks and formatting more uniformly
    full_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    # Regex pattern to match user and AI markers with optional whitespace after the marker
    pattern = re.compile(r"(##USER_MESSAGE##:|##AI_RESPONSE##:|Entry \d+:)\s*")

    # Split text based on the markers while retaining the markers for context
    split_entries = pattern.split(full_text)

    current_marker = None

    for part in split_entries:
        part = part.strip()
        if not part:
            continue

        # Detect the marker and set the current context (user or AI)
        if re.match(r"##USER_MESSAGE##:", part):
            if current_entry["user_message"] or current_entry["ai_response"]:
                entries.append(current_entry)
                current_entry = {"user_message": "", "ai_response": ""}

            current_marker = "user"
        
        elif re.match(r"##AI_RESPONSE##:", part):
            current_marker = "ai"

        elif re.match(r"Entry \d+:", part):
            # If encountering "Entry <number>", store the previous entry and start a new one
            if current_entry["user_message"] or current_entry["ai_response"]:
                entries.append(current_entry)
                current_entry = {"user_message": "", "ai_response": ""}
            # Reset current_marker as it does not indicate user or AI content but marks a new entry
            current_marker = None

        # Accumulate text based on the current marker
        else:
            if current_marker == "user":
                current_entry["user_message"] += part + " "
            elif current_marker == "ai":
                current_entry["ai_response"] += part + " "

    # Add the last entry if it wasn't already added
    if current_entry["user_message"] or current_entry["ai_response"]:
        entries.append(current_entry)

    # Process each entry and store it in the database
    entry_number = 1
    for entry in entries:
        user_msg = entry["user_message"].strip()
        ai_resp = entry["ai_response"].strip()

        if user_msg or ai_resp:
            combined_text = user_msg + " " + ai_resp
            tags_str = ", ".join(process_text_for_nltk(combined_text, custom_stop_words))
            people_names = get_continuous_chunks(combined_text)
            people_names_str = ", ".join(people_names)

            # Store the entry in the database with an additional entry_number field
            insert_into_db_ingestion(
                user_id=user_id,
                entry_date=entry_date,
                entry_number=entry_number,  # Set the entry number explicitly here
                user_message=user_msg,
                ai_response=ai_resp,
                daily_people_names=people_names_str,
                tags_str=tags_str
            )
            print(f">>> [ Daily_file_ingestion.py - ingest_word_doc ] - DB Ingested entry {entry_number} for [{entry_date}] to user [{user_id}]")
            entry_number += 1




# def ingest_word_doc(user_id, file_path, entry_date):
#     doc = Document(file_path)
#     user_message = ""
#     ai_response = ""
#     entry_number = 1  # Initialize entry number
#     entries = []

#     # Parsing logic (this part iterates over paragraphs and splits entries by markers)
#     for paragraph in doc.paragraphs:
#         text = paragraph.text.strip()

#         if text.startswith("##USER_MESSAGE##:"):
#             if user_message or ai_response:
#                 # Store the current entry before starting a new one
#                 entries.append((user_message.strip(), ai_response.strip(), entry_number))
#                 user_message = ""
#                 ai_response = ""
#                 entry_number += 1  # Increment entry number

#             user_message = text[5:].strip() + " "
        
#         elif text.startswith("##AI_RESPONSE##:"):
#             ai_response = text.split(":", 1)[1].strip() if ":" in text else text.strip() + " "

#         else:
#             # Continue accumulating text
#             if "##USER_MESSAGE##:" in text:
#                 user_message += text + " "
#             elif "##AI_RESPONSE##" in text:
#                 ai_response += text + " "

#     # Add the last entry if it wasn't already added
#     if user_message or ai_response:
#         entries.append((user_message.strip(), ai_response.strip(), entry_number))

#     # Store each entry in the database
#     for user_msg, ai_resp, entry_num in entries:
#         insert_into_db_ingestion(
#             user_id=user_id,
#             entry_date=entry_date,
#             entry_number=entry_num,  # Pass entry_number
#             user_message=user_msg,
#             ai_response=ai_resp,
#             daily_people_names="",
#             tags_str=""
#         )



# def ingest_word_doc(user_id, file_path, entry_date):
#     doc = Document(file_path)
#     user_message = ""
#     ai_response = ""
#     entry_number = 1
#     entries = []

#     # Join all paragraphs together to handle line breaks and formatting more uniformly
#     full_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

#     # Regex pattern to match the user and AI response markers
#     pattern = re.compile(r"(##USER_MESSAGE##:|##AI_RESPONSE##:)")

#     # Split text based on the markers while retaining the markers for context
#     split_entries = pattern.split(full_text)

#     current_marker = None

#     for part in split_entries:
#         part = part.strip()
#         if not part:
#             continue

#         # Detect the marker and set the current context (user or AI)
#         if part == "##USER_MESSAGE##:":
#             if user_message or ai_response:
#                 # Store the previous entry before starting a new one
#                 entries.append((user_message.strip(), ai_response.strip(), entry_number))
#                 user_message = ""
#                 ai_response = ""
#                 entry_number += 1

#             current_marker = "USER"

#         elif part == "##AI_RESPONSE##:":
#             if user_message or ai_response:
#                 # Store the previous entry before starting a new one
#                 entries.append((user_message.strip(), ai_response.strip(), entry_number))
#                 user_message = ""
#                 ai_response = ""
#                 entry_number += 1

#             current_marker = "AI"

#         # Accumulate text based on the current marker
#         else:
#             if current_marker == "USER":
#                 user_message += part + " "
#             elif current_marker == "AI":
#                 ai_response += part + " "

#     # Add the last entry if it wasn't already added
#     if user_message or ai_response:
#         entries.append((user_message.strip(), ai_response.strip(), entry_number))

#     # Process each entry and store it in the database
#     for user_msg, ai_resp, entry_num in entries:
#         combined_text = user_msg + " " + ai_resp
#         tags_str = ", ".join(process_text_for_nltk(combined_text, custom_stop_words))
#         people_names = get_continuous_chunks(combined_text)
#         people_names_str = ", ".join(people_names)

#         # Create a unique identifier for each entry (for logging purposes)
#         unique_entry_id = f"{entry_date} - {entry_num}"

#         # Store the entry in the database
#         insert_into_db_ingestion(
#             user_id=user_id,
#             entry_date=entry_date,
#             entry_number=entry_num,
#             user_message=user_msg,
#             ai_response=ai_resp,
#             daily_people_names=people_names_str,
#             tags_str=tags_str
#         )
#         print(f">>> [ File Ingestion ] - Stored entry {unique_entry_id} for user {user_id}")




# # # # # # 
# Chat GPT Routes and Endpoints
# RESPONSE
# DAILY DIARY
# # # # # # 
def process_chatgpt_response(entry_date, user_message, ai_response):
    # Combine user_message and ai_response for analysis
    combined_text = user_message + " " + ai_response

    # Tokenize the text using WordPunctTokenizer
    tokenizer = WordPunctTokenizer()
    tokens = tokenizer.tokenize(combined_text.lower())

    # Remove stop words and short words
    stop_words = set(stopwords.words('english'))
    stop_words.update(custom_stop_words) # Add your custom stop words to stop_words.json
    filtered_tokens = [word for word in tokens if word not in stop_words and word.isalpha() and len(word) > 2]

    # Part-of-speech tagging
    pos_tags = nltk.pos_tag(filtered_tokens)

    # Keep only nouns and adjectives
    meaningful_words = [word for word, pos in pos_tags if pos in ('NN', 'JJ', 'VB', 'NNP', 'NNS', 'VBG')]

    # Frequency analysis
    word_counts = Counter(meaningful_words)
    common_words = [word for word, count in word_counts.most_common(50)]  # Top 50 common words

    return common_words